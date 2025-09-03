#!/usr/bin/env python3
import argparse
import re
from datetime import datetime
from pathlib import Path
import pandas as pd
import numpy as np

def parse_sheet_date(name: str):
    name = name.strip()
    m = re.match(r"^(\d{1,2})[/-](\d{1,2})[/-](\d{4})$", name)
    if m:
        d, mth, y = map(int, m.groups())
        return pd.Timestamp(year=y, month=mth, day=d).date()
    return None

def money(x):
    try:
        return f"${int(round(float(x), 0)):,}".replace(",", ".")
    except Exception:
        return str(x)

def fmt_date(d):
    return pd.to_datetime(d).strftime("%d/%m/%Y")

def load_ledger(excel_path: Path):
    xls = pd.ExcelFile(excel_path)
    records = []
    for s in xls.sheet_names:
        try:
            df = pd.read_excel(excel_path, sheet_name=s, header=0)
        except Exception:
            continue
        df.columns = [str(c).strip().upper() for c in df.columns]
        col_map = {}
        for c in df.columns:
            if "CUENTA" in c:
                col_map["CUENTA"] = c
            elif "DEBE" in c:
                col_map["DEBE"] = c
            elif "HABER" in c:
                col_map["HABER"] = c
        if "CUENTA" not in col_map or "DEBE" not in col_map:
            base_cols = df.columns[:3].tolist()
            while len(base_cols) < 3:
                base_cols.append(None)
            col_map = {"CUENTA": base_cols[0], "DEBE": base_cols[1], "HABER": base_cols[2] if len(base_cols) > 2 else None}
        cuenta_col = col_map.get("CUENTA")
        debe_col = col_map.get("DEBE")
        haber_col = col_map.get("HABER")
        temp = pd.DataFrame({
            "Fecha": parse_sheet_date(s),
            "Cuenta": df[cuenta_col] if cuenta_col in df.columns else pd.Series(dtype=object),
            "Debe": pd.to_numeric(df[debe_col], errors="coerce") if debe_col in df.columns else pd.Series(dtype=float),
            "Haber": pd.to_numeric(df[haber_col], errors="coerce") if (haber_col and haber_col in df.columns) else pd.Series(dtype=float)
        })
        temp = temp[~(temp["Cuenta"].isna() & temp["Debe"].isna() & temp["Haber"].isna())].copy()
        temp = temp[temp["Cuenta"].astype(str).str.upper() != "CUENTA"]
        temp["Cuenta"] = temp["Cuenta"].astype(str).str.strip()
        records.append(temp)
    all_df = pd.concat(records, ignore_index=True) if records else pd.DataFrame(columns=["Fecha","Cuenta","Debe","Haber"])
    all_df = all_df[~all_df["Fecha"].isna()].copy()
    all_df["Fecha"] = pd.to_datetime(all_df["Fecha"]).dt.date
    return all_df

def compute_sections(all_df: pd.DataFrame, petty_cash_name: str, min_transfer_debe: float, transfer_keywords: list[str]):
    PETTY_CASH_KEY = petty_cash_name.upper()
    retiros = (all_df[(all_df["Cuenta"].str.upper() == PETTY_CASH_KEY) & (all_df["Debe"] > 0)]
               .groupby("Fecha", as_index=False)["Debe"].sum()
               .rename(columns={"Debe":"Monto"}))
    retiros["Concepto"] = "COMPRAS"
    detalles = all_df[(all_df["Debe"] > 0) & (all_df["Cuenta"].str.upper() != PETTY_CASH_KEY)].copy()
    det_totals = detalles.groupby("Fecha", as_index=False)["Debe"].sum().rename(columns={"Debe":"Detalle_Total"})
    retiros = retiros.merge(det_totals, on="Fecha", how="left")
    retiros["Detalle_Total"] = retiros["Detalle_Total"].fillna(0)
    retiros["Efectivo"] = (retiros["Monto"] - retiros["Detalle_Total"]).round(0)

    if all_df["Haber"].notna().any() and all_df["Haber"].fillna(0).sum() > 0:
        posibles_trans = all_df[all_df["Haber"] > 0].copy()
        posibles_trans["Tipo"] = "HABER"
        valor_col = "Haber"
    else:
        mask_kw = all_df["Cuenta"].str.upper().apply(lambda x: any(k in x for k in transfer_keywords))
        posibles_trans = all_df[(all_df["Debe"] >= min_transfer_debe) & mask_kw].copy()
        posibles_trans["Tipo"] = "DEBE(MAYOR)"
        valor_col = "Debe"
    trans_diario = posibles_trans.groupby("Fecha", as_index=False)[valor_col].sum().rename(columns={valor_col:"Monto_Transferido"})
    total_trans = float(trans_diario["Monto_Transferido"].sum()) if not trans_diario.empty else 0.0

    return retiros, detalles, posibles_trans, trans_diario, total_trans

def build_docx(out_path: Path, empresa: str, mes: str, responsable: str,
               retiros: pd.DataFrame, detalles: pd.DataFrame,
               posibles_trans: pd.DataFrame, trans_diario: pd.DataFrame, total_trans: float):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except Exception as e:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

    doc = Document()
    doc.add_heading(f"Empresa: {empresa}", level=1)
    doc.add_paragraph(f"Mes: {mes}  •  Responsable: {responsable or '(pendiente)'}  •  Fecha de elaboración: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_heading("DETALLE DE RETIROS DIARIOS PARA CAJA", level=2)
    if not retiros.empty:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        hdr[0].text = "Fecha"
        hdr[1].text = "Monto retirado"
        hdr[2].text = "Concepto / Observación"
        for _, r in retiros.sort_values("Fecha").iterrows():
            row = table.add_row().cells
            row[0].text = fmt_date(r["Fecha"])
            row[1].text = money(r["Monto"])
            row[2].text = str(r["Concepto"])
        p = doc.add_paragraph(f"TOTAL CAJAS: {money(retiros['Monto'].sum())}")
        p.runs[0].bold = True
    else:
        doc.add_paragraph("No se detectaron retiros de 'CAJA MENOR'. Verifique el nombre exacto de la cuenta.")

    doc.add_heading("DETALLES DE CAJAS MENORES POR DÍA", level=2)
    if not detalles.empty:
        for fecha, sub in detalles.groupby("Fecha"):
            doc.add_heading(fmt_date(fecha), level=3)
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = table.rows[0].cells
            hdr[0].text = "Concepto"
            hdr[1].text = "Valor"
            for _, row in sub.sort_values("Cuenta").iterrows():
                r = table.add_row().cells
                r[0].text = str(row["Cuenta"])
                r[1].text = money(row["Debe"])
            tot = sub["Debe"].sum()
            fila_retiro = retiros[retiros["Fecha"] == fecha]
            if not fila_retiro.empty:
                retiro_val = float(fila_retiro["Monto"].values[0])
                efectivo = float(fila_retiro["Efectivo"].values[0])
                p = doc.add_paragraph(f"TOTAL: {money(tot)}    EFECTIVO: {money(efectivo)}")
                p.runs[0].bold = True
                doc.add_paragraph(
                    f"El día {fmt_date(fecha)} se realiza un retiro para compras por un valor de {money(retiro_val)} "
                    f"el cual en facturación tiene un valor de {money(tot)} y el mensajero entrega en efectivo {money(efectivo)}."
                )
            doc.add_paragraph("")
    else:
        doc.add_paragraph("No se detectaron egresos de caja (DEBE) distintos a 'CAJA MENOR'.")

    doc.add_heading("DETALLES DE TRANSFERENCIAS", level=2)
    if not trans_diario.empty:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Fecha"
        hdr[1].text = "Monto transferido"
        hdr[2].text = "Concepto / Observación"
        for _, r in trans_diario.sort_values("Fecha").iterrows():
            row = table.add_row().cells
            row[0].text = fmt_date(r["Fecha"])
            row[1].text = money(r["Monto_Transferido"])
            row[2].text = "PROVEEDORES / GASTOS (estimado)"
        p = doc.add_paragraph(f"TOTAL: {money(total_trans)}")
        p.runs[0].bold = True

        doc.add_heading("2.2 DETALLES DE TRANSFERENCIAS POR DÍA", level=3)
        for fecha, sub in posibles_trans.groupby("Fecha"):
            doc.add_paragraph(fmt_date(fecha)).runs[0].bold = True
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text = "Concepto"
            hdr[1].text = "Valor"
            for _, row in sub.sort_values("Cuenta").iterrows():
                val = row["Haber"] if (not pd.isna(row["Haber"]) and row["Haber"] > 0) else row["Debe"]
                r = table.add_row().cells
                r[0].text = str(row["Cuenta"])
                r[1].text = money(val)
            doc.add_paragraph(f"TOTAL: {money(sub['Debe'].sum())}")
            doc.add_paragraph("")
    else:
        doc.add_paragraph("No se detectaron transferencias de manera inequívoca. Si su libro registra transferencias en otra hoja o columna, indíquelo y ajusto el criterio.")

    doc.save(out_path)

def main():
    ap = argparse.ArgumentParser(description="Construye un informe .docx desde un libro diario en Excel.")
    ap.add_argument("--input", required=True, help="Ruta del Excel con varias hojas por día.")
    ap.add_argument("--out", default="informe_movimientos.docx", help="Ruta del archivo .docx a generar.")
    ap.add_argument("--empresa", default="Surtiprocesos industriales S.A.S")
    ap.add_argument("--mes", default="Agosto de 2025")
    ap.add_argument("--responsable", default="(pendiente)")
    ap.add_argument("--petty-cash-name", default="CAJA MENOR")
    ap.add_argument("--min-transfer-debe", type=float, default=300000)
    ap.add_argument("--transfer-keywords", default="PROVEEDOR,PROVEEDORES,BANCO,ARRIENDO,SERVICIO,SERVICIOS,PLANILLA,VEHICULO,VEHICULOS,GASOLINA,CLARO,COMPENSAR,FALABELLA,FINANDINA,DOTACION")
    args = ap.parse_args()

    excel_path = Path(args.input)
    out_path = Path(args.out)
    if not excel_path.exists():
        raise SystemExit(f"[ERROR] No se encuentra el archivo: {excel_path}")

    all_df = load_ledger(excel_path)
    keywords = [k.strip().upper() for k in args.transfer_keywords.split(",") if k.strip()]
    retiros, detalles, posibles_trans, trans_diario, total_trans = compute_sections(
        all_df,
        petty_cash_name=args.petty_cash_name,
        min_transfer_debe=args.min_transfer_debe,
        transfer_keywords=keywords
    )
    try:
        build_docx(out_path, args.empresa, args.mes, args.responsable,
                   retiros, detalles, posibles_trans, trans_diario, total_trans)
    except RuntimeError as e:
        print(str(e))
        raise
    print(f"[OK] Informe creado: {out_path.resolve()}")

if __name__ == "__main__":
    main()
